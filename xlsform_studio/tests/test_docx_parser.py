"""Tests for the Word (.docx) parser (Module 1 / Iteration 4)."""

import docx
import pytest

from xlsform_studio.parsers.docx_parser import DocxParser


@pytest.fixture()
def sample_docx(tmp_path):
    d = docx.Document()
    d.add_heading("SECTION A: REGISTRATION", level=1)
    d.add_paragraph("Are you a returning customer?")
    d.add_paragraph("Yes")
    d.add_paragraph("No")
    d.add_paragraph("If yes, record membership date.")
    d.add_paragraph("Number of previous orders")
    d.add_heading("SECTION B: DETAILS", level=1)
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Preferred contact method"
    table.rows[0].cells[1].text = "Email / Phone"
    d.add_paragraph("Total amount spent")
    path = tmp_path / "survey.docx"
    d.save(path)
    return path


def test_docx_question_count(sample_docx):
    qn = DocxParser().parse(sample_docx)
    labels = [q.raw_label for q in qn.questions]
    assert "Number of previous orders" in labels
    assert "Total amount spent" in labels
    assert len(qn.questions) == 4


def test_docx_yes_no_options(sample_docx):
    qn = DocxParser().parse(sample_docx)
    first = qn.questions[0]
    assert first.raw_choices == ["Yes", "No"]
    assert "membership date" in first.logic


def test_docx_table_options(sample_docx):
    qn = DocxParser().parse(sample_docx)
    contact = next(q for q in qn.questions if q.raw_label == "Preferred contact method")
    assert contact.raw_choices == ["Email", "Phone"]


def test_docx_sections(sample_docx):
    qn = DocxParser().parse(sample_docx)
    sections = {q.section for q in qn.questions}
    assert "Registration" in sections
    assert "Details" in sections


def test_missing_file_raises():
    with pytest.raises(FileNotFoundError):
        DocxParser().parse("/no/such/file.docx")
