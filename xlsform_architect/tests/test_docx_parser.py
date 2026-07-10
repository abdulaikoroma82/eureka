"""Tests for the Word (.docx) parser (Module 1 / Iteration 4)."""

import docx
import pytest

from xlsform_architect.parsers.docx_parser import DocxParser


@pytest.fixture()
def sample_docx(tmp_path):
    d = docx.Document()
    d.add_heading("SECTION A: SCREENING", level=1)
    d.add_paragraph("Is the child currently enrolled in OTP?")
    d.add_paragraph("Yes")
    d.add_paragraph("No")
    d.add_paragraph("If yes, record admission date.")
    d.add_paragraph("Child age in months")
    d.add_heading("SECTION B: ANTHROPOMETRY", level=1)
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Child sex"
    table.rows[0].cells[1].text = "Male / Female"
    d.add_paragraph("MUAC (cm)")
    path = tmp_path / "survey.docx"
    d.save(path)
    return path


def test_docx_question_count(sample_docx):
    qn = DocxParser().parse(sample_docx)
    labels = [q.raw_label for q in qn.questions]
    assert "Child age in months" in labels
    assert "MUAC (cm)" in labels
    assert len(qn.questions) == 4


def test_docx_yes_no_options(sample_docx):
    qn = DocxParser().parse(sample_docx)
    otp = qn.questions[0]
    assert otp.raw_choices == ["Yes", "No"]
    assert "admission date" in otp.logic


def test_docx_table_options(sample_docx):
    qn = DocxParser().parse(sample_docx)
    sex = next(q for q in qn.questions if q.raw_label == "Child sex")
    assert sex.raw_choices == ["Male", "Female"]


def test_docx_sections(sample_docx):
    qn = DocxParser().parse(sample_docx)
    sections = {q.section for q in qn.questions}
    assert "Screening" in sections
    assert "Anthropometry" in sections


def test_missing_file_raises():
    with pytest.raises(FileNotFoundError):
        DocxParser().parse("/no/such/file.docx")
