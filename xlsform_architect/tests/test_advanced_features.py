"""Tests for the gap-closing features: compound logic, grids, repeats,
other-specify, list dedup, cascading/translation passthrough, PDF noise."""

import io

import docx
import openpyxl
import pytest

from xlsform_architect.app.workflow import Workflow
from xlsform_architect.engine.logic_engine import LogicEngine
from xlsform_architect.models import Question, Questionnaire
from xlsform_architect.parsers.docx_parser import DocxParser
from xlsform_architect.parsers.pdf_parser import PdfParser
from xlsform_architect.parsers.questionnaire_parser import QuestionnaireParser
from xlsform_architect.xlsform.survey_builder import SurveyBuilder


def _run(data, **kw):
    kw.setdefault("write_outputs", False)
    return Workflow().run_from_dict(data, **kw)


# --- Module 5: compound & comparison logic -----------------------------------
def test_compound_and_condition():
    r = _run({"settings": {"form_title": "T", "form_id": "t"}, "survey": [
        {"question": "Are you employed?", "choices": ["Yes", "No"]},
        {"question": "Respondent age"},
        {"question": "Occupation", "logic": "if yes and age over 18"}]})
    occ = next(q for q in r.questionnaire.questions if q.name == "occupation")
    assert occ.relevant == "${employed}='1' and ${respondent_age}>18"


def test_compound_or_condition():
    r = _run({"settings": {"form_title": "T", "form_id": "t"}, "survey": [
        {"question": "Respondent age"},
        {"question": "Consent form", "logic": "if age under 18 or age over 65"}]})
    consent = next(q for q in r.questionnaire.questions if q.name == "consent_form")
    assert consent.relevant == "${respondent_age}<18 or ${respondent_age}>65"


def test_at_least_comparison():
    r = _run({"settings": {"form_title": "T", "form_id": "t"}, "survey": [
        {"question": "Household size (number of members)"},
        {"question": "Crowding question", "logic": "if household size is at least 5"}]})
    q = next(x for x in r.questionnaire.questions if x.name == "crowding_question")
    assert q.relevant.endswith(">=5")


def test_selected_for_multi_select_equality():
    r = _run({"settings": {"form_title": "T", "form_id": "t"}, "survey": [
        {"question": "Select all services used", "choices": ["Clinic", "Pharmacy"]},
        {"question": "Clinic details", "logic": "if services is clinic"}]})
    q = next(x for x in r.questionnaire.questions if x.name == "clinic_details")
    assert q.relevant.startswith("selected(")


def test_skip_to_gets_honest_note_not_guess():
    q = Question(raw_label="Q5", name="q5", logic="If no, skip to question 20")
    LogicEngine().resolve(q, previous=None, known=[q])
    assert q.relevant == ""
    assert any("Skip pattern" in a for a in q.assumptions)


def test_bare_yes_binds_to_nearest_yes_no():
    r = _run({"settings": {"form_title": "T", "form_id": "t"}, "survey": [
        {"question": "Do you own the dwelling?", "choices": ["Yes", "No"]},
        {"question": "Number of rooms"},
        {"question": "Purchase year", "logic": "if yes"}]})
    q = next(x for x in r.questionnaire.questions if x.name == "purchase_year")
    assert q.relevant == "${own_dwelling}='1'"


# --- matrix / grid questions ---------------------------------------------------
@pytest.fixture()
def grid_docx(tmp_path):
    d = docx.Document()
    d.add_paragraph("Please rate the following:")
    t = d.add_table(rows=4, cols=4)
    for i, h in enumerate(["", "Poor", "Fair", "Good"]):
        t.rows[0].cells[i].text = h
    for r, item in enumerate(["Cleanliness", "Staff friendliness",
                              "Waiting time"], start=1):
        t.rows[r].cells[0].text = item
    path = tmp_path / "grid.docx"
    d.save(path)
    return path


def test_grid_becomes_selects_with_shared_list(grid_docx):
    qn = DocxParser().parse(grid_docx)
    result = Workflow().run(qn, form_title="G", form_id="g", write_outputs=False)
    assert result.is_valid
    selects = [q for q in result.questionnaire.questions if q.is_select]
    assert len(selects) == 3
    # All three share ONE deduplicated list.
    lists = {q.xlsform_type for q in selects}
    assert len(lists) == 1


def test_two_column_table_still_question_options(tmp_path):
    d = docx.Document()
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Preferred contact"
    t.rows[0].cells[1].text = "Email / Phone"
    path = tmp_path / "twocol.docx"
    d.save(path)
    qn = DocxParser().parse(path)
    assert qn.questions[0].raw_choices == ["Email", "Phone"]


# --- repeat groups ----------------------------------------------------------------
def test_json_repeat_flag_emits_begin_repeat():
    r = _run({"settings": {"form_title": "T", "form_id": "t"}, "survey": [
        {"question": "Member name", "section": "Members", "repeat": True},
        {"question": "Member age", "section": "Members", "repeat": True}]})
    rows = SurveyBuilder().build(r.questionnaire)
    types = [x["type"] for x in rows]
    assert "begin repeat" in types and "end repeat" in types
    assert r.is_valid


def test_text_for_each_heading_becomes_repeat():
    qn = QuestionnaireParser().parse_text(
        "FOR EACH HOUSEHOLD MEMBER\nMember name\nMember age in years\n")
    assert all(q.section_type == "repeat" for q in qn.questions)
    result = Workflow().run(qn, form_title="T", form_id="t", write_outputs=False)
    rows = SurveyBuilder().build(result.questionnaire)
    assert any(x["type"] == "begin repeat" for x in rows)
    assert result.is_valid


def test_explicit_structural_rows_pass_through():
    r = _run({"settings": {"form_title": "T", "form_id": "t"}, "survey": [
        {"type": "begin repeat", "name": "hh_member", "label": "Household member"},
        {"question": "Member name"},
        {"type": "end repeat", "name": "hh_member"}]})
    rows = SurveyBuilder().build(r.questionnaire)
    assert rows[0]["type"] == "begin repeat"
    assert rows[-1]["type"] == "end repeat"
    assert r.is_valid


# --- Other (specify) -----------------------------------------------------------------
def test_other_specify_followup_injected():
    r = _run({"settings": {"form_title": "T", "form_id": "t"}, "survey": [
        {"question": "Main water source",
         "choices": ["Piped", "Well", "Other (specify)"]}]})
    follow = next((q for q in r.questionnaire.questions
                   if q.name.endswith("_other")), None)
    assert follow is not None
    assert follow.xlsform_type == "text"
    assert follow.relevant.startswith("selected(${main_water_source}")
    assert r.is_valid


def test_no_other_no_followup():
    r = _run({"settings": {"form_title": "T", "form_id": "t"}, "survey": [
        {"question": "Gender", "choices": ["Male", "Female"]}]})
    assert not any(q.name.endswith("_other") for q in r.questionnaire.questions)


# --- choice list dedup ------------------------------------------------------------------
def test_identical_likert_lists_shared():
    r = _run({"settings": {"form_title": "T", "form_id": "t"}, "survey": [
        {"question": "Rate the food", "choices": ["Poor", "Fair", "Good"]},
        {"question": "Rate the service", "choices": ["Poor", "Fair", "Good"]},
        {"question": "Rate the price", "choices": ["Poor", "Fair", "Good"]}]})
    non_yesno = [k for k in r.questionnaire.choice_lists if k != "yes_no"]
    assert len(non_yesno) == 1
    selects = [q for q in r.questionnaire.questions if q.is_select]
    assert len({q.xlsform_type for q in selects}) == 1


def test_different_lists_not_merged():
    r = _run({"settings": {"form_title": "T", "form_id": "t"}, "survey": [
        {"question": "Rate the food", "choices": ["Poor", "Good"]},
        {"question": "Meal time", "choices": ["Morning", "Evening"]}]})
    non_yesno = [k for k in r.questionnaire.choice_lists if k != "yes_no"]
    assert len(non_yesno) == 2


# --- cascading selects & translations ------------------------------------------------------
def test_choice_filter_and_translations_exported():
    r = _run({
        "settings": {"form_title": "T", "form_id": "t"},
        "choices": {
            "region": [{"name": "north", "label": "North",
                        "label::French (fr)": "Nord"}],
            "district": [{"name": "d1", "label": "District 1",
                          "region": "north"}],
        },
        "survey": [
            {"question": "Region", "type": "select_one region",
             "label::French (fr)": "Région"},
            {"question": "District", "type": "select_one district",
             "choice_filter": "region=${region}"},
        ]})
    wb = openpyxl.load_workbook(io.BytesIO(r.xlsform_bytes))
    survey_headers = [c.value for c in wb["survey"][1] if c.value]
    choices_headers = [c.value for c in wb["choices"][1] if c.value]
    assert "choice_filter" in survey_headers
    assert "label::French (fr)" in survey_headers
    assert "label::French (fr)" in choices_headers
    assert "region" in choices_headers          # cascade filter column
    assert r.is_valid


# --- PDF noise filtering ---------------------------------------------------------------------
def test_pdf_page_numbers_and_running_headers_dropped(tmp_path):
    import fitz
    doc = fitz.open()
    for i in range(1, 4):
        page = doc.new_page()
        page.insert_text((72, 72), "ACME Survey 2026")           # running header
        page.insert_text((72, 120), f"Question {i}: How old are you?")
        page.insert_text((72, 700), f"Page {i} of 3")             # footer
    path = tmp_path / "noisy.pdf"
    doc.save(str(path))
    doc.close()

    lines = PdfParser().extract_lines(path)
    assert not any("Page" in ln for ln in lines)
    assert "ACME Survey 2026" not in lines
    assert any("How old are you" in ln for ln in lines)


# --- messy real-world document stress test ------------------------------------------------------
def test_messy_docx_end_to_end(tmp_path):
    d = docx.Document()
    d.add_heading("COMMUNITY SERVICES SURVEY", level=0)
    d.add_paragraph("Confidential - for internal use only")
    d.add_heading("SECTION A: RESPONDENT", level=1)
    d.add_paragraph("1. What is your full name?")
    d.add_paragraph("2. Are you a resident of this district?")
    d.add_paragraph("Yes")
    d.add_paragraph("No")
    d.add_paragraph("If yes, record the number of years of residence.")
    d.add_heading("SECTION B: SERVICE RATINGS", level=1)
    t = d.add_table(rows=3, cols=4)
    for i, h in enumerate(["Service", "Poor", "OK", "Good"]):
        t.rows[0].cells[i].text = h
    t.rows[1].cells[0].text = "Water supply"
    t.rows[2].cells[0].text = "Road maintenance"
    d.add_heading("FOR EACH CHILD IN THE HOUSEHOLD", level=1)
    d.add_paragraph("Child name")
    d.add_paragraph("Child date of birth")
    d.add_paragraph("Any other comments?")
    path = tmp_path / "messy.docx"
    d.save(path)

    result = Workflow().run_from_file(path, form_title="Messy",
                                      form_id="messy", write_outputs=False)
    assert result.is_valid, [f.message for f in result.report.errors]
    qn = result.questionnaire
    names = [q.name for q in qn.questions]
    # Yes/No question compiled with follow-up logic
    resident = next(q for q in qn.questions if "resident" in q.name)
    assert resident.xlsform_type == "select_one yes_no"
    # Grid rows became selects sharing a list
    grid_qs = [q for q in qn.questions
               if q.is_select and q.name != resident.name]
    assert len(grid_qs) == 2
    assert len({q.xlsform_type for q in grid_qs}) == 1
    # Repeat roster detected
    rows = SurveyBuilder().build(qn)
    assert any(r["type"] == "begin repeat" for r in rows)
    # DOB produced an age calculation
    assert "age_years" in names
